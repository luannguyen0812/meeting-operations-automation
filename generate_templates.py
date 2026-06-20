#!/usr/bin/env python3
"""Generate Meeting Agenda and Meeting Minutes .docx templates for all projects."""

import sys
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config import PROJECTS, AGENDA_ROWS

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'output')


def calculate_dates(target_date):
    """Calculate the meeting date and next meeting date."""
    weekday = target_date.weekday()  # 0=Mon, 4=Fri
    if weekday == 0:  # Monday
        next_meeting = target_date + timedelta(days=4)  # Friday
    elif weekday == 4:  # Friday
        next_meeting = target_date + timedelta(days=3)  # Next Monday
    else:
        next_meeting = target_date + timedelta(days=1)
    return target_date, next_meeting


def format_date_filename(d):
    return d.strftime('%m.%d.%Y')


def format_date_header(d):
    return d.strftime('%A, %B %d, %Y')


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return run


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '000000',
        })
        borders.append(el)
    tblPr.append(borders)


def create_agenda(project, meeting_date, next_meeting_date):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Header
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{project["name"]}')
    run.font.color.rgb = RGBColor(0, 51, 153)
    run.font.size = Pt(18)

    subtitle = doc.add_heading(level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Meeting Agenda')
    run.font.color.rgb = RGBColor(0, 51, 153)
    run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(format_date_header(meeting_date))
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Meeting Details table
    doc.add_heading('Meeting Details', level=3)
    details = [
        ('Date', format_date_header(meeting_date)),
        ('Time', project['time']),
        ('Duration', '30 minutes'),
        ('Facilitator', project['facilitator']),
        ('Attendees', ', '.join(project['attendees'])),
        ('Apologies', ''),
        ('Location/Link', project['meet_link']),
    ]
    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    for i, (label, value) in enumerate(details):
        set_cell_text(table.cell(i, 0), label, bold=True)
        set_cell_text(table.cell(i, 1), value)

    doc.add_paragraph()

    # Objective
    doc.add_heading('Objective', level=3)
    doc.add_paragraph('To review project progress, address challenges, and align on upcoming priorities.')

    # Pre-Read
    doc.add_heading('Pre-Read / Preparation', level=3)
    doc.add_paragraph('')

    # Agenda table
    doc.add_heading('Agenda', level=3)
    agenda_table = doc.add_table(rows=len(AGENDA_ROWS) + 1, cols=5)
    agenda_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(agenda_table)

    headers = ['#', 'Topic', 'Owner', 'Time', 'Notes']
    for i, h in enumerate(headers):
        set_cell_text(agenda_table.cell(0, i), h, bold=True)

    for row_idx, (num, topic, owner, time) in enumerate(AGENDA_ROWS):
        set_cell_text(agenda_table.cell(row_idx + 1, 0), num)
        set_cell_text(agenda_table.cell(row_idx + 1, 1), topic)
        set_cell_text(agenda_table.cell(row_idx + 1, 2), owner)
        set_cell_text(agenda_table.cell(row_idx + 1, 3), time)
        set_cell_text(agenda_table.cell(row_idx + 1, 4), '')

    doc.add_paragraph()

    # Next Meeting
    doc.add_heading('Next Meeting', level=3)
    doc.add_paragraph(f'Date: {format_date_header(next_meeting_date)}')
    doc.add_paragraph(f'Time: {project["time"]}')
    doc.add_paragraph('Tentative Agenda: To be confirmed')

    return doc


def create_minutes(project, meeting_date, next_meeting_date):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Header
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{project["name"]}')
    run.font.color.rgb = RGBColor(0, 51, 153)
    run.font.size = Pt(18)

    subtitle = doc.add_heading(level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Meeting Minutes')
    run.font.color.rgb = RGBColor(0, 51, 153)
    run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(format_date_header(meeting_date))
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Meeting Details table
    doc.add_heading('Meeting Details', level=3)
    details = [
        ('Date', format_date_header(meeting_date)),
        ('Time', project['time']),
        ('Duration', '30 minutes'),
        ('Facilitator', project['facilitator']),
        ('Attendees', ', '.join(project['attendees'])),
        ('Apologies', ''),
        ('Location/Link', project['meet_link']),
    ]
    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    for i, (label, value) in enumerate(details):
        set_cell_text(table.cell(i, 0), label, bold=True)
        set_cell_text(table.cell(i, 1), value)

    doc.add_paragraph()

    # Discussion Summary table
    doc.add_heading('Discussion Summary', level=3)
    disc_table = doc.add_table(rows=7, cols=4)
    disc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(disc_table)

    disc_headers = ['#', 'Topic Discussed', 'Key Points/Decisions', 'Action Items']
    for i, h in enumerate(disc_headers):
        set_cell_text(disc_table.cell(0, i), h, bold=True)
    for r in range(1, 7):
        set_cell_text(disc_table.cell(r, 0), str(r))
        for c in range(1, 4):
            set_cell_text(disc_table.cell(r, c), '')

    doc.add_paragraph()

    # Action Items table
    doc.add_heading('Action Items', level=3)
    action_table = doc.add_table(rows=5, cols=4)
    action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(action_table)

    action_headers = ['Action Item', 'Owner', 'Due Date', 'Status']
    for i, h in enumerate(action_headers):
        set_cell_text(action_table.cell(0, i), h, bold=True)
    for r in range(1, 5):
        for c in range(4):
            set_cell_text(action_table.cell(r, c), '')

    doc.add_paragraph()

    # Next Steps
    doc.add_heading('Next Steps / Follow-up', level=3)
    doc.add_paragraph('')

    # Additional Notes
    doc.add_heading('Additional Notes', level=3)
    doc.add_paragraph('')

    return doc


def main():
    if len(sys.argv) > 1:
        target_date = datetime.strptime(sys.argv[1], '%Y-%m-%d').date()
    else:
        target_date = (datetime.now().date() + timedelta(days=1))

    meeting_date, next_meeting_date = calculate_dates(target_date)
    date_str = format_date_filename(meeting_date)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    generated = []

    for project in PROJECTS:
        agenda_filename = f'Meeting Agenda {date_str}.docx'
        minutes_filename = f'Meeting Minutes {date_str}.docx'

        project_dir = os.path.join(OUTPUT_DIR, project['short_name'])
        os.makedirs(project_dir, exist_ok=True)

        agenda_doc = create_agenda(project, meeting_date, next_meeting_date)
        agenda_path = os.path.join(project_dir, agenda_filename)
        agenda_doc.save(agenda_path)

        minutes_doc = create_minutes(project, meeting_date, next_meeting_date)
        minutes_path = os.path.join(project_dir, minutes_filename)
        minutes_doc.save(minutes_path)

        generated.append({
            'project': project['short_name'],
            'agenda': agenda_path,
            'minutes': minutes_path,
        })
        print(f'✓ {project["short_name"]}: {agenda_filename}, {minutes_filename}')

    print(f'\nGenerated {len(generated) * 2} files for {format_date_header(meeting_date)}')
    print(f'Next meeting: {format_date_header(next_meeting_date)}')
    return generated


if __name__ == '__main__':
    main()
