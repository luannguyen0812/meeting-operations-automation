#!/usr/bin/env python3
"""
Workflow 2: Post-Meeting Minutes Distribution
Pulls HappyScribe transcript → AI summary → downloads Drive template → fills it →
PDF → email (PDF attachment) → upload PDF to Drive → cleanup local temp files.
"""

import os
import sys
import json
import shutil
import tempfile
import urllib.request
import urllib.parse
from datetime import datetime, date, timedelta

from config import PROJECTS
from roster import correct_transcript_names
from happyscribe import find_todays_transcripts, export_transcript_text, match_transcript_to_project, get_transcription
from send_email import send_minutes_email

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

_SPARKY_TOKEN_FILE = os.path.expanduser('~/.openclaw/telegram-sparky.token')
_LUAN_CHAT_ID = '8318978943'


def _notify_luan(text):
    try:
        token = open(_SPARKY_TOKEN_FILE).read().strip()
        url = f'https://api.telegram.org/bot{token}/sendMessage'
        payload = urllib.parse.urlencode({'chat_id': _LUAN_CHAT_ID, 'text': text}).encode()
        urllib.request.urlopen(url, payload, timeout=10)
    except Exception as e:
        print(f'Telegram notify failed (non-fatal): {e}')


def find_project_by_name(name):
    name_lower = name.lower()
    for p in PROJECTS:
        if name_lower in p['name'].lower() or name_lower in p['short_name'].lower():
            return p
    return None


def get_drive_service():
    import json as _json
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from config import GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET
    token_path = os.path.join(SCRIPT_DIR, 'token.json')
    with open(token_path) as f:
        td = _json.load(f)
    creds = Credentials(
        token=td['token'], refresh_token=td['refresh_token'],
        token_uri=td['token_uri'], client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
    )
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        td['token'] = creds.token
        with open(token_path, 'w') as f:
            _json.dump(td, f, indent=2)
    return build('drive', 'v3', credentials=creds)


def download_drive_minutes_template(service, project, date_str):
    """Download the blank Meeting Minutes .docx from Drive. Returns local temp path or None."""
    import io
    from googleapiclient.http import MediaIoBaseDownload

    folder_id = project['drive_folder_id']
    filename = f'Meeting Minutes {date_str}.docx'

    query = f"'{folder_id}' in parents and name='{filename}' and trashed=false"
    results = service.files().list(q=query, fields='files(id, name)').execute()
    files = results.get('files', [])
    if not files:
        return None

    file_id = files[0]['id']
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    request = service.files().get_media(fileId=file_id)
    downloader = MediaIoBaseDownload(tmp, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    tmp.close()
    return tmp.name


def get_apologies(transcript_obj, project):
    """Cross-reference HappyScribe participants against project attendee list."""
    participants = transcript_obj.get('meeting', {}).get('participants', [])
    present_names = {p.get('name', '').lower() for p in participants if p.get('name')}

    apologies = []
    for attendee in project.get('attendees', []):
        first = attendee.split()[0].lower()
        last = attendee.split()[-1].lower() if len(attendee.split()) > 1 else ''
        if not any(first in p or last in p for p in present_names):
            apologies.append(attendee)

    return ', '.join(apologies) if apologies else ''


def ai_summarize(transcript_text, project_name, roster=None):
    """Use the local claude CLI to produce a structured meeting summary."""
    import subprocess

    roster_section = ''
    if roster:
        names = ', '.join(m['full_name'] for m in roster)
        roster_section = f'\nKnown team members (use these exact names — the transcript may have misspellings or phonetic errors): {names}\n'

    prompt = f"""You are summarizing a recorded meeting for {project_name}.
{roster_section}
Transcript:
{transcript_text}

Return a JSON object with exactly these keys:
- "topics": list of objects {{"topic": str, "key_points": str, "action_items": str}} — up to 6 topics discussed
- "action_items": list of objects {{"item": str, "owner": str, "due_date": str, "status": str}} — up to 5 concrete action items
- "next_steps": str — 1-2 sentences on overall next steps
- "additional_notes": str — anything else worth noting, or empty string

Rules:
- Be concise and factual. No invented details.
- owner should be a person's name if mentioned, else "Team"
- due_date should be a specific date if mentioned, else "TBD"
- status is always "Pending" unless explicitly stated otherwise
- When assigning an owner, match to the closest known team member name above.
- Return only the JSON, no markdown fences."""

    result = subprocess.run(
        ['/usr/local/bin/claude', '--print', '--tools', ''],
        input=prompt,
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        # claude CLI writes auth/login errors to stdout, not stderr — include both
        detail = (result.stderr.strip() or result.stdout.strip() or f'exit {result.returncode}')
        raise RuntimeError(f'claude CLI error: {detail}')
    raw = result.stdout.strip()
    if raw.startswith('```'):
        raw = raw.split('\n', 1)[1].rsplit('```', 1)[0].strip()
    return json.loads(raw)


def fill_existing_docx(docx_path, project, summary, meeting_date, apologies):
    """Fill blank cells in the downloaded Drive template with AI summary data."""
    from docx import Document
    from docx.shared import Pt

    def set_cell(cell, text):
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    def find_paragraph_after_heading(doc, heading_text):
        for i, para in enumerate(doc.paragraphs):
            if heading_text.lower() in para.text.lower() and para.style.name.startswith('Heading'):
                if i + 1 < len(doc.paragraphs):
                    return doc.paragraphs[i + 1]
        return None

    doc = Document(docx_path)
    tables = doc.tables

    # Table 0: Meeting Details — fill Apologies row (index 5)
    if tables and len(tables[0].rows) > 5:
        set_cell(tables[0].cell(5, 1), apologies)

    # Table 1: Discussion Summary — fill rows 1..N
    if len(tables) > 1:
        topics = summary.get('topics', [])
        disc_table = tables[1]
        for i, topic in enumerate(topics):
            row_idx = i + 1
            if row_idx >= len(disc_table.rows):
                disc_table.add_row()
                set_cell(disc_table.cell(row_idx, 0), str(row_idx))
            set_cell(disc_table.cell(row_idx, 1), topic.get('topic', ''))
            set_cell(disc_table.cell(row_idx, 2), topic.get('key_points', ''))
            set_cell(disc_table.cell(row_idx, 3), topic.get('action_items', ''))

    # Table 2: Action Items — fill rows 1..N
    if len(tables) > 2:
        actions = summary.get('action_items', [])
        act_table = tables[2]
        for i, item in enumerate(actions):
            row_idx = i + 1
            if row_idx >= len(act_table.rows):
                act_table.add_row()
            set_cell(act_table.cell(row_idx, 0), item.get('item', ''))
            set_cell(act_table.cell(row_idx, 1), item.get('owner', ''))
            set_cell(act_table.cell(row_idx, 2), item.get('due_date', 'TBD'))
            set_cell(act_table.cell(row_idx, 3), item.get('status', 'Pending'))

    # Next Steps and Additional Notes free-text paragraphs
    next_para = find_paragraph_after_heading(doc, 'Next Steps')
    if next_para:
        next_para.clear()
        next_para.add_run(summary.get('next_steps', ''))

    notes_para = find_paragraph_after_heading(doc, 'Additional Notes')
    if notes_para:
        notes_para.clear()
        notes_para.add_run(summary.get('additional_notes', ''))

    doc.save(docx_path)


def build_minutes_docx_from_scratch(project, summary, meeting_date, apologies):
    """Fallback: generate a filled .docx when no Drive template exists."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    def set_cell(cell, text, bold=False, size=10):
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(text))
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        run.bold = bold

    def add_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        borders = tblPr.makeelement(qn('w:tblBorders'), {})
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = borders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'single', qn('w:sz'): '4',
                qn('w:space'): '0', qn('w:color'): '000000',
            })
            borders.append(el)
        tblPr.append(borders)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(project['name'])
    r.font.color.rgb = RGBColor(0, 51, 153)
    r.font.size = Pt(18)

    sub = doc.add_heading(level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Meeting Minutes')
    r.font.color.rgb = RGBColor(0, 51, 153)
    r.font.size = Pt(14)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run(meeting_date.strftime('%A, %B %d, %Y'))
    r.font.color.rgb = RGBColor(0, 102, 204)
    r.font.size = Pt(12)
    doc.add_paragraph()

    doc.add_heading('Meeting Details', level=3)
    details = [
        ('Date', meeting_date.strftime('%A, %B %d, %Y')),
        ('Time', project.get('time', '')),
        ('Duration', '30 minutes'),
        ('Facilitator', project.get('facilitator', 'Luan Nguyen')),
        ('Attendees', ', '.join(project.get('attendees', []))),
        ('Apologies', apologies),
        ('Location/Link', project.get('meet_link', '')),
    ]
    t = doc.add_table(rows=len(details), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(t)
    for i, (label, val) in enumerate(details):
        set_cell(t.cell(i, 0), label, bold=True)
        set_cell(t.cell(i, 1), val)
    doc.add_paragraph()

    doc.add_heading('Discussion Summary', level=3)
    topics = summary.get('topics', [])
    disc = doc.add_table(rows=max(len(topics), 1) + 1, cols=4)
    disc.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(disc)
    for i, h in enumerate(['#', 'Topic Discussed', 'Key Points/Decisions', 'Action Items']):
        set_cell(disc.cell(0, i), h, bold=True)
    for i, topic in enumerate(topics, 1):
        set_cell(disc.cell(i, 0), str(i))
        set_cell(disc.cell(i, 1), topic.get('topic', ''))
        set_cell(disc.cell(i, 2), topic.get('key_points', ''))
        set_cell(disc.cell(i, 3), topic.get('action_items', ''))
    doc.add_paragraph()

    doc.add_heading('Action Items', level=3)
    actions = summary.get('action_items', [])
    act = doc.add_table(rows=max(len(actions), 1) + 1, cols=4)
    act.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(act)
    for i, h in enumerate(['Action Item', 'Owner', 'Due Date', 'Status']):
        set_cell(act.cell(0, i), h, bold=True)
    for i, item in enumerate(actions, 1):
        set_cell(act.cell(i, 0), item.get('item', ''))
        set_cell(act.cell(i, 1), item.get('owner', ''))
        set_cell(act.cell(i, 2), item.get('due_date', 'TBD'))
        set_cell(act.cell(i, 3), item.get('status', 'Pending'))
    doc.add_paragraph()

    doc.add_heading('Next Steps / Follow-up', level=3)
    doc.add_paragraph(summary.get('next_steps', ''))
    doc.add_heading('Additional Notes', level=3)
    doc.add_paragraph(summary.get('additional_notes', ''))

    return doc


def docx_to_pdf(docx_path):
    from docx2pdf import convert
    pdf_path = docx_path.replace('.docx', '.pdf')
    convert(docx_path, pdf_path)
    return pdf_path


def update_drive_docx(service, local_docx_path, project, date_str):
    """Replace the blank Drive template with the filled .docx (in place)."""
    from googleapiclient.http import MediaFileUpload

    folder_id = project['drive_folder_id']
    filename = f'Meeting Minutes {date_str}.docx'

    # Find existing file ID
    query = f"'{folder_id}' in parents and name='{filename}' and trashed=false"
    results = service.files().list(q=query, fields='files(id)').execute()
    files = results.get('files', [])

    media = MediaFileUpload(
        local_docx_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=True,
    )

    if files:
        updated = service.files().update(
            fileId=files[0]['id'],
            media_body=media,
            fields='id',
        ).execute()
        return updated['id']
    else:
        created = service.files().create(
            body={'name': filename, 'parents': [folder_id]},
            media_body=media,
            fields='id',
        ).execute()
        return created['id']


def export_docx_as_pdf_via_drive(service, local_docx_path):
    """Convert a local .docx to PDF bytes.

    Drive's export endpoint only works on native Docs Editors files, not on
    uploaded .docx files. So we import the .docx as a temporary Google Doc
    (Drive converts it on the way in), export that as PDF, then delete it.
    Runs headless — no Word, no LibreOffice, no macOS permission prompts.
    """
    import io
    from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload

    media = MediaFileUpload(
        local_docx_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=True,
    )
    temp_doc = service.files().create(
        body={
            'name': 'tmp-minutes-pdf-export',
            'mimeType': 'application/vnd.google-apps.document',
        },
        media_body=media,
        fields='id',
    ).execute()
    temp_id = temp_doc['id']
    try:
        request = service.files().export_media(fileId=temp_id, mimeType='application/pdf')
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue()
    finally:
        service.files().delete(fileId=temp_id).execute()


def process_meeting(project_name, transcript_id=None, meeting_date=None):
    """Full post-meeting pipeline."""
    if meeting_date is None:
        meeting_date = date.today()

    date_str = meeting_date.strftime('%m.%d.%Y')

    project = find_project_by_name(project_name)
    if not project:
        print(f'Project not found: {project_name}')
        return False

    print(f'\n=== Processing: {project["name"]} ===\n')

    # 1. Get transcript text + raw object (for participants)
    if transcript_id:
        print(f'Using transcript: {transcript_id}')
        transcript_obj = get_transcription(transcript_id)
        transcript_text = export_transcript_text(transcript_id)
    else:
        transcripts = find_todays_transcripts(meeting_date)
        project_transcripts = [t for t in transcripts if match_transcript_to_project(t, [project])]
        if not project_transcripts:
            print(f'No transcript found for {project["name"]} on {meeting_date}')
            return False
        transcript_obj = project_transcripts[-1]
        transcript_id = transcript_obj['id']
        print(f'Found transcript: {transcript_obj.get("name", "?")}')
        transcript_text = export_transcript_text(transcript_id)

    # 2. Correct names
    print('Correcting names...')
    from roster import get_master_roster, download_workbook, get_project_members
    wb = download_workbook()
    project_members = get_project_members(project['name'], wb)
    master = get_master_roster(wb)
    all_members = {m['full_name']: m for m in master}
    for pm in project_members:
        all_members.setdefault(pm['full_name'], pm)
    roster = list(all_members.values())
    corrected_text, _ = correct_transcript_names(transcript_text, project['name'])

    # 3. Derive apologies from HappyScribe participant list
    apologies = get_apologies(transcript_obj, project)
    if apologies:
        print(f'Apologies: {apologies}')
    else:
        print('Apologies: none detected')

    # 4. AI summary
    print('Generating AI summary...')
    summary = ai_summarize(corrected_text, project['name'], roster=roster)

    # 5. Get Drive service; try to download existing template
    print('Fetching Drive template...')
    tmp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(tmp_dir, f'Meeting Minutes {date_str}.docx')

    try:
        service = get_drive_service()
        drive_template = download_drive_minutes_template(service, project, date_str)

        if drive_template:
            print('  Found Drive template — filling in blanks.')
            shutil.copy(drive_template, docx_path)
            os.unlink(drive_template)
            fill_existing_docx(docx_path, project, summary, meeting_date, apologies)
        else:
            print('  No Drive template found — generating from scratch.')
            doc = build_minutes_docx_from_scratch(project, summary, meeting_date, apologies)
            doc.save(docx_path)

        # 6. Push filled .docx back to Drive (replaces blank template in place)
        print('Updating Drive minutes template with filled content...')
        file_id = update_drive_docx(service, docx_path, project, date_str)
        print(f'  Drive file updated: {file_id}')

        # 7. Export PDF via a throwaway Google Doc conversion (Drive can't
        #    export uploaded .docx directly — only native Docs Editors files)
        print('Exporting PDF (via temp Google Doc conversion)...')
        pdf_bytes = export_docx_as_pdf_via_drive(service, docx_path)
        pdf_filename = f'Meeting Minutes {date_str}.pdf'

        # 8. Send email with PDF attachment
        body = (
            f'Hi team,\n\nPlease find attached the meeting minutes for the '
            f'{project["name"]} session held on {meeting_date.strftime("%A, %B %d, %Y")}.\n\n'
            f'Best regards,\nLuan Nguyen'
        )
        print(f'Sending minutes to {len(project["emails"])} recipients...')
        success = send_minutes_email(
            to_emails=project['emails'],
            project_name=project['name'],
            date_str=date_str,
            body_text=body,
            attachment_bytes=pdf_bytes,
            attachment_filename=pdf_filename,
        )
        if not success:
            print('Email failed.')
            return False

        _notify_luan(f'✓ {project["name"]} minutes sent to {len(project["emails"])} recipients')
        return True

    finally:
        # 9. Always clean up local temp files
        shutil.rmtree(tmp_dir, ignore_errors=True)
        print('Local temp files cleaned up.')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python workflow_postmeeting.py "Project Name" [transcript_id] [YYYY-MM-DD]')
        sys.exit(1)

    p = sys.argv[1]
    tid = sys.argv[2] if len(sys.argv) > 2 else None
    mdate = datetime.strptime(sys.argv[3], '%Y-%m-%d').date() if len(sys.argv) > 3 else None
    process_meeting(p, tid, mdate)
